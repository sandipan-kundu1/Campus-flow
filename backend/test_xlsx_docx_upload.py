import sys
import os
import io
import json
import openpyxl
import docx

sys.path.insert(0, os.path.dirname(__file__))

from app.utils.text_extractor import extract_text
from app.routes.upload import _parse_timetable_text_with_llm

def create_mock_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Timetable"
    # Write some headers and data
    headers = ["Day", "Subject", "Start Time", "End Time", "Room", "Instructor"]
    ws.append(headers)
    
    rows = [
        ["Monday", "Database Systems (DBMS)", "09:00 AM", "10:30 AM", "Room 401", "Dr. Sandipan"],
        ["Tuesday", "Computer Networks (CN)", "11:00 AM", "12:30 PM", "Lab 2", "Prof. Joy"],
        ["Wednesday", "Web Development", "02:00 PM", "03:30 PM", "Seminar Hall", "Dr. Roy"],
        ["Thursday", "Compiler Design", "09:00 AM", "10:30 AM", "Room 403", "Prof. Roy"],
        ["Friday", "DBMS Lab", "11:00 AM", "01:00 PM", "Lab 4", "Dr. Sandipan"]
    ]
    for row in rows:
        ws.append(row)
        
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()

def create_mock_docx():
    doc = docx.Document()
    doc.add_heading("Weekly Class Timetable", level=1)
    
    table = doc.add_table(rows=6, cols=6)
    headers = ["Day", "Subject", "Time", "End Time", "Room", "Instructor"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        
    rows = [
        ["Monday", "Database Systems (DBMS)", "09:00 AM", "10:30 AM", "Room 401", "Dr. Sandipan"],
        ["Tuesday", "Computer Networks (CN)", "11:00 AM", "12:30 PM", "Lab 2", "Prof. Joy"],
        ["Wednesday", "Web Development", "02:00 PM", "03:30 PM", "Seminar Hall", "Dr. Roy"],
        ["Thursday", "Compiler Design", "09:00 AM", "10:30 AM", "Room 403", "Prof. Roy"],
        ["Friday", "DBMS Lab", "11:00 AM", "01:00 PM", "Lab 4", "Dr. Sandipan"]
    ]
    
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def run_test():
    print("Generating mock XLSX...")
    xlsx_bytes = create_mock_xlsx()
    print("Extracting XLSX text...")
    xlsx_text = extract_text(xlsx_bytes, "timetable.xlsx")
    print(f"Extracted XLSX Text:\n{xlsx_text}\n")
    
    print("Parsing XLSX with LLM...")
    xlsx_entries = _parse_timetable_text_with_llm(xlsx_text)
    print(f"Parsed XLSX Entries:\n{json.dumps(xlsx_entries, indent=2)}\n")

    print("Generating mock DOCX...")
    docx_bytes = create_mock_docx()
    print("Extracting DOCX text...")
    docx_text = extract_text(docx_bytes, "timetable.docx")
    print(f"Extracted DOCX Text:\n{docx_text}\n")
    
    print("Parsing DOCX with LLM...")
    docx_entries = _parse_timetable_text_with_llm(docx_text)
    print(f"Parsed DOCX Entries:\n{json.dumps(docx_entries, indent=2)}\n")

if __name__ == "__main__":
    run_test()
