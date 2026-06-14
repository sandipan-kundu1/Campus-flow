import openpyxl
import docx

def create_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Timetable"
    headers = ["Day", "Subject", "Start Time", "End Time", "Room", "Instructor"]
    ws.append(headers)
    rows = [
        ["Monday", "Database Systems (DBMS)", "09:00 AM", "10:30 AM", "Room 401", "Dr. Sandipan"],
        ["Tuesday", "Computer Networks (CN)", "11:00 AM", "12:30 PM", "Lab 2", "Prof. Joy"],
        ["Wednesday", "Web Development", "02:00 PM", "03:30 PM", "Seminar Hall", "Dr. Roy"],
        ["Thursday", "Compiler Design", "09:00 AM", "10:30 AM", "Room 403", "Prof. Roy"],
        ["Friday", "DBMS Lab", "11:00 AM", "01:00 PM", "Lab 4", "Dr. Sandipan"]
    ]
    for row in rows:
        ws.append(row)
    wb.save("test_timetable.xlsx")
    print("Created test_timetable.xlsx")

def create_docx():
    doc = docx.Document()
    doc.add_heading("Weekly Class Timetable", level=1)
    table = doc.add_table(rows=6, cols=6)
    headers = ["Day", "Subject", "Time", "End Time", "Room", "Instructor"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
    rows = [
        ["Monday", "Database Systems (DBMS)", "09:00 AM", "10:30 AM", "Room 401", "Dr. Sandipan"],
        ["Tuesday", "Computer Networks (CN)", "11:00 AM", "12:30 PM", "Lab 2", "Prof. Joy"],
        ["Wednesday", "Web Development", "02:00 PM", "03:30 PM", "Seminar Hall", "Dr. Roy"],
        ["Thursday", "Compiler Design", "09:00 AM", "10:30 AM", "Room 403", "Prof. Roy"],
        ["Friday", "DBMS Lab", "11:00 AM", "01:00 PM", "Lab 4", "Dr. Sandipan"]
    ]
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
    doc.save("test_timetable.docx")
    print("Created test_timetable.docx")

if __name__ == "__main__":
    create_xlsx()
    create_docx()
